from pathlib import Path
import re

from docx import Document as DocxDocument
from pdfminer.high_level import extract_text


class ParsingError(Exception):
    pass


def extract_text_from_file(file_path: str) -> tuple[str, dict]:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = extract_text(file_path)
        if not text or not text.strip():
            raise ParsingError("OCR_TODO")
        metadata = _build_metadata(text)
        metadata["pages"] = text.count("\f") + 1
        return text, metadata
    if suffix == ".docx":
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        metadata = _build_metadata(text)
        metadata["paragraphs"] = len(doc.paragraphs)
        return text, metadata
    raise ParsingError("UNSUPPORTED_FORMAT")


def _build_metadata(text: str) -> dict:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return {
        "skills": _extract_skills(text),
        "experience": _extract_section(lines, ("experience", "work history", "employment")),
        "education": _extract_section(lines, ("education", "academics", "studies")),
    }


def _extract_skills(text: str) -> list[str]:
    skills_catalog = {
        "python",
        "javascript",
        "typescript",
        "react",
        "next.js",
        "node",
        "fastapi",
        "docker",
        "kubernetes",
        "postgresql",
        "redis",
        "aws",
        "gcp",
        "azure",
        "sql",
        "mongodb",
        "graphql",
        "tailwind",
        "figma",
    }
    found = set()
    lowered = text.lower()
    for skill in skills_catalog:
        pattern = rf"\\b{re.escape(skill)}\\b"
        if re.search(pattern, lowered):
            found.add(skill)
    return sorted(found)


def _extract_section(lines: list[str], headings: tuple[str, ...]) -> list[str]:
    heading_index = None
    for index, line in enumerate(lines):
        if any(keyword in line.lower() for keyword in headings):
            heading_index = index
            break
    if heading_index is None:
        return _fallback_section(lines, headings)
    section = []
    for line in lines[heading_index + 1 :]:
        if _looks_like_heading(line):
            break
        section.append(line)
        if len(section) >= 5:
            break
    return section


def _fallback_section(lines: list[str], headings: tuple[str, ...]) -> list[str]:
    keywords = {
        "education": ("university", "college", "bachelor", "master", "phd", "b.sc", "m.sc"),
        "experience": ("engineer", "developer", "manager", "analyst", "designer", "intern"),
    }
    key = "education" if "education" in headings[0] else "experience"
    matches = []
    for line in lines:
        if any(word in line.lower() for word in keywords[key]):
            matches.append(line)
        if len(matches) >= 5:
            break
    return matches


def _looks_like_heading(line: str) -> bool:
    lowered = line.lower()
    if lowered in {"skills", "experience", "work history", "education"}:
        return True
    return line.isupper()
